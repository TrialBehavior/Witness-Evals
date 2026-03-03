from flask import Flask, render_template, request, jsonify, send_file
import pyreadstat
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import io
import base64
from datetime import datetime
import os
from werkzeug.utils import secure_filename
from config import *

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_CONFIG['upload_folder']
app.config['MAX_CONTENT_LENGTH'] = UPLOAD_CONFIG['max_file_size_mb'] * 1024 * 1024
app.config['SECRET_KEY'] = 'your-secret-key-here'

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Store session data (in production, use proper session management)
session_data = {}

def parse_characteristic_name(variable_label):
    """Extract the characteristic name from the SPSS variable label"""
    delimiter = CHARACTERISTIC_PARSER['delimiter']
    
    if delimiter in variable_label:
        parts = variable_label.split(delimiter)
        if len(parts) > 1:
            # Get the last word after the delimiter
            characteristic = parts[-1].strip().split()[-1]
            return characteristic
    return variable_label

def calculate_distribution_percentages(series):
    """Calculate distribution percentages that sum to 100%"""
    value_counts = series.value_counts()
    total = len(series)
    
    # Calculate raw percentages
    raw_percentages = {val: (count / total) * 100 for val, count in value_counts.items()}
    
    # Round to nearest integer
    rounded = {val: round(pct) for val, pct in raw_percentages.items()}
    
    # Adjust to ensure sum is 100
    current_sum = sum(rounded.values())
    diff = 100 - current_sum
    
    if diff != 0:
        # Sort by decimal part to determine which values to adjust
        decimals = [(val, raw_percentages[val] - rounded[val]) for val in rounded.keys()]
        decimals.sort(key=lambda x: x[1], reverse=(diff > 0))
        
        # Adjust the values with largest decimal parts
        for i in range(abs(diff)):
            val = decimals[i][0]
            rounded[val] += 1 if diff > 0 else -1
    
    return rounded

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not file.filename.lower().endswith('.sav'):
        return jsonify({'error': 'Please upload an SPSS (.sav) file'}), 400
    
    try:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Read SPSS file
        df, meta = pyreadstat.read_sav(filepath)
        
        # Store in session
        session_id = str(hash(filename + str(datetime.now())))
        session_data[session_id] = {
            'filepath': filepath,
            'df': df,
            'meta': meta
        }
        
        # Get variable names and labels
        variables = []
        for var in df.columns:
            label = meta.column_names_to_labels.get(var, var)
            variables.append({
                'name': var,
                'label': label
            })
        
        return jsonify({
            'success': True,
            'session_id': session_id,
            'variables': variables,
            'n_rows': len(df)
        })
    
    except Exception as e:
        return jsonify({'error': f'Error reading file: {str(e)}'}), 500

@app.route('/generate', methods=['POST'])
def generate_document():
    try:
        data = request.json
        session_id = data['session_id']
        witness_name = data['witness_name']
        date_str = data['date']
        task1_vars = data['task1_variables']
        task2_vars = data['task2_variables']  # Now expects a list of 1 or 2 variables
        case_name = data['case_name']
        case_id = data['case_id']
        footer_date_str = data['footer_date']
        witness_number = data['witness_number']
        
        # Retrieve session data
        if session_id not in session_data:
            return jsonify({'error': 'Session expired'}), 400
        
        df = session_data[session_id]['df']
        meta = session_data[session_id]['meta']
        
        # Create Word document
        doc = Document()
        
        # Set narrow margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(DOCUMENT_CONFIG['margins']['top'])
            section.bottom_margin = Inches(DOCUMENT_CONFIG['margins']['bottom'])
            section.left_margin = Inches(DOCUMENT_CONFIG['margins']['left'])
            section.right_margin = Inches(DOCUMENT_CONFIG['margins']['right'])
        
        # Date (top left)
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        date_formatted = date_obj.strftime(DATE_FORMAT)
        date_para = doc.add_paragraph(date_formatted)
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Witness Evaluation (centered) - REDUCED SPACING
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(f'Witness Evaluation: {witness_name}')
        title_run.font.size = Pt(DOCUMENT_CONFIG['title_font_size'])
        title_run.font.bold = True
        # Reduce space after title
        title_para.paragraph_format.space_after = Pt(3)
        
        # Sample size - REDUCED SPACING with italic n
        n = len(df)
        n_para = doc.add_paragraph()
        n_para.add_run('(')
        n_run = n_para.add_run('n')
        n_run.italic = True
        n_para.add_run(f' = {n})')
        n_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Reduce space after n
        n_para.paragraph_format.space_after = Pt(3)
        
        # Add horizontal line
        add_horizontal_line(doc)
        
        # Task 1: Grid
        task1_para = doc.add_paragraph()
        # Split the template and insert witness name properly
        task1_text = TEXT_TEMPLATES['task1_intro']
        parts = task1_text.split('{witness_name}')
        
        task1_para.add_run('1. ' + parts[0])
        task1_para.add_run(witness_name).bold = True
        if len(parts) > 1:
            task1_para.add_run(parts[1])
        
        # Create grid table
        create_characteristics_table(doc, df, meta, task1_vars)
        
        # Triple space
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Task 2: Graph(s) - Handle 1 or 2 variables
        # Ensure task2_vars is a list
        if not isinstance(task2_vars, list):
            task2_vars = [task2_vars]
        
        num_graphs = len(task2_vars)
        
        if num_graphs == 1:
            # Single graph - original behavior
            task2_var = task2_vars[0]
            
            task2_para = doc.add_paragraph()
            task2_text = TEXT_TEMPLATES['task2_intro'].format(witness_name='{witness_name}')
            # Split and handle bold witness name
            parts = task2_text.split('{witness_name}')
            task2_para.add_run('2. ' + parts[0])
            task2_para.add_run(witness_name).bold = True
            if len(parts) > 1:
                task2_para.add_run(parts[1])
            
            # Generate and add graph
            graph_stream = create_side_graph(df, meta, task2_var)
            doc.add_picture(graph_stream, width=Inches(6))
            
        else:
            # Two graphs - one for each defendant
            for idx, task2_var in enumerate(task2_vars):
                # Get the defendant name from the variable label
                var_label = meta.column_names_to_labels.get(task2_var, task2_var)
                defendant_name = extract_defendant_name(var_label)
                
                task2_para = doc.add_paragraph()
                # Format: "2. Which side did Bob Joe's testimony help the most? (vs. Defendant Name)"
                if idx == 0:
                    task2_text = TEXT_TEMPLATES['task2_intro'].format(witness_name='{witness_name}')
                    parts = task2_text.split('{witness_name}')
                    task2_para.add_run('2. ' + parts[0])
                    task2_para.add_run(witness_name).bold = True
                    if len(parts) > 1:
                        task2_para.add_run(parts[1])
                    
                else:
                    # Second graph - same question but different defendant
                    task2_text = TEXT_TEMPLATES['task2_intro'].format(witness_name='{witness_name}')
                    parts = task2_text.split('{witness_name}')
                    task2_para.add_run('   ' + parts[0])  # Indent slightly
                    task2_para.add_run(witness_name).bold = True
                    if len(parts) > 1:
                        task2_para.add_run(parts[1])

                
                # Generate and add graph
                graph_stream = create_side_graph(df, meta, task2_var)
                doc.add_picture(graph_stream, width=Inches(6))
                
                # Add space between graphs if not the last one
                if idx < len(task2_vars) - 1:
                    doc.add_paragraph()
        
        # Add footer to all pages
        add_footer(doc, case_name, case_id, footer_date_str, witness_number)
        
        # Save document
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f'witness_eval_{session_id}.docx')
        doc.save(output_path)
        
        return send_file(output_path, 
                        as_attachment=True, 
                        download_name=f'Witness_Evaluation_{witness_name.replace(" ", "_")}.docx',
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    except Exception as e:
        return jsonify({'error': f'Error generating document: {str(e)}'}), 500

def extract_defendant_name(variable_label):
    """Extract defendant name from variable label for multi-defendant cases"""
    # Try to find text after common patterns like "vs." or "v."
    import re
    
    # Look for patterns like "vs. John Doe" or "v. John Doe"
    patterns = [
        r'vs\.?\s+(.+?)(?:\s+testimony|\s+help|\?|$)',
        r'v\.?\s+(.+?)(?:\s+testimony|\s+help|\?|$)',
        r'against\s+(.+?)(?:\s+testimony|\s+help|\?|$)',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, variable_label, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    
    # Fallback: return "Defendant" if no name found
    return "Defendant"

def add_horizontal_line(doc):
    """Add a horizontal line to the document"""
    para = doc.add_paragraph()
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)

def create_characteristics_table(doc, df, meta, task1_vars):
    """Create the characteristics grid table"""
    labels = list(VALUE_LABELS.values())
    
    # Create table
    n_rows = len(task1_vars) + 1  # +1 for header
    table = doc.add_table(rows=n_rows, cols=7)
    table.style = DOCUMENT_CONFIG['table_style']
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(TABLE_HEADERS):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Increase header font size
        for run in header_cells[i].paragraphs[0].runs:
            run.font.size = Pt(12)
    
    # Data rows
    for idx, var in enumerate(task1_vars):
        row = table.rows[idx + 1]
        
        # Get variable label and parse characteristic
        var_label = meta.column_names_to_labels.get(var, var)
        characteristic = parse_characteristic_name(var_label)
        
        # Calculate mean
        mean_val = df[var].mean()
        row.cells[0].text = f'{mean_val:.2f}'
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Characteristic name
        row.cells[1].text = characteristic
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Distribution percentages
        dist_pct = calculate_distribution_percentages(df[var])
        
        for i, val in enumerate([1, 2, 3, 4, 5], start=2):
            pct = dist_pct.get(val, 0)
            row.cells[i].text = f'{pct}%'
            row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Apply font size and cell padding to all cells in the row
        for cell in row.cells:
            # Increase font size
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(13)  # Adjust this value (try 12-14)
            
            # Add cell padding to make rows taller
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            
            # Add padding to top and bottom
            for margin in ['top', 'bottom']:
                node = OxmlElement(f'w:{margin}')
                node.set(qn('w:w'), '100')  # Padding in twips (120 = about 0.08 inches)
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            
            tcPr.append(tcMar)

def create_side_graph(df, meta, task2_var):
    """Create bar graph for which side testimony helped"""
    import textwrap
    
    # Get value labels
    var_label = meta.column_names_to_labels.get(task2_var, task2_var)
    value_labels = meta.variable_value_labels.get(task2_var, {})
    
    # Calculate distribution
    dist_pct = calculate_distribution_percentages(df[task2_var])
    
    # Prepare data for plotting
    categories = []
    percentages = []
    colors = []
    
    for val in sorted(dist_pct.keys()):
        label = value_labels.get(val, str(val))
        categories.append(label)
        percentages.append(dist_pct[val])
        
        # Determine color based on keywords
        label_lower = label.lower()
        color = CHART_CONFIG['colors']['other']  # default
        
        for keyword in SIDE_KEYWORDS['plaintiff']:
            if keyword.lower() in label_lower:
                color = CHART_CONFIG['colors']['plaintiff']
                break
        
        for keyword in SIDE_KEYWORDS['defendant']:
            if keyword.lower() in label_lower:
                color = CHART_CONFIG['colors']['defendant']
                break
        
        colors.append(color)
    
    # Create plot with improved styling
    fig, ax = plt.subplots(figsize=CHART_CONFIG['figure_size'])
    
    # Create bars with shadow effect
    bars = ax.bar(range(len(categories)), percentages, color=colors, 
                   edgecolor='black', linewidth=1.5, alpha=0.85)
    
    # Add subtle gradient/texture effect using hatching
    for bar, color in zip(bars, colors):
        bar.set_linewidth(1.5)
        bar.set_edgecolor('black')
        # Add subtle shadow by drawing a darker bar slightly offset
        height = bar.get_height()
        x = bar.get_x()
        width = bar.get_width()
        shadow = plt.Rectangle((x + 0.02, 0.05), width, height - 0.05, 
                               facecolor='gray', alpha=0.3, zorder=0)
        ax.add_patch(shadow)
    
    # Add percentage labels on top of bars
    for i, (bar, pct) in enumerate(zip(bars, percentages)):
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height,
                f'{pct}%',
                ha='center', va='bottom', 
                fontsize=CHART_CONFIG['font_sizes']['percentage'], 
                fontweight='bold')
    
    # Formatting
    ax.set_ylabel('Percentage', fontsize=CHART_CONFIG['font_sizes']['ylabel'], fontweight='bold')
    ax.set_xticks(range(len(categories)))
    
    # Wrap text to fit within bar width - approximately 12 characters per line for bar width
    wrapped_labels = []
    for cat in categories:
        # Use textwrap to wrap at ~12 characters per line (adjusts based on bar width)
        wrapped = textwrap.fill(cat, width=12)
        wrapped_labels.append(wrapped)
    
    ax.set_xticklabels(wrapped_labels, rotation=0, ha='center', 
                       fontsize=CHART_CONFIG['font_sizes']['labels'], fontweight='bold')
    ax.set_ylim(0, max(percentages) * 1.15)
    
    # Add % symbol to y-axis tick labels
    from matplotlib.ticker import FuncFormatter
    ax.yaxis.set_major_formatter(FuncFormatter(lambda y, _: f'{int(y)}%'))
    
    # Grid for easier reading
    ax.yaxis.grid(True, linestyle='--', alpha=0.3, zorder=0)
    ax.set_axisbelow(True)
    
    # Style spines
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_linewidth(1.5)
    ax.spines['bottom'].set_linewidth(1.5)
    
    plt.tight_layout()
    
    # Save to bytes
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=CHART_CONFIG['dpi'], bbox_inches='tight')
    img_stream.seek(0)
    plt.close()
    
    return img_stream

def add_footer(doc, case_name, case_id, footer_date_str, witness_number):
    """Add footer with case info, page number, and confidentiality notice"""
    from datetime import datetime
    
    # Format the footer date
    footer_date_obj = datetime.strptime(footer_date_str, '%Y-%m-%d')
    footer_date_formatted = footer_date_obj.strftime(DATE_FORMAT)
    
    # Access the default section (all pages will have the same footer)
    section = doc.sections[0]
    footer = section.footer
    
    # Clear any existing footer content
    footer.paragraphs[0].clear()
    
    # Create a table with 3 columns for footer layout
    table = footer.add_table(rows=2, cols=3, width=Inches(7))
    
    # Set column widths
    table.columns[0].width = Inches(2.5)  # Left
    table.columns[1].width = Inches(2)    # Center
    table.columns[2].width = Inches(2.5)  # Right
    
    # Row 1: Case info (left), Page number (center), Confidentiality (right)
    # Left cell - Case name and ID
    left_cell = table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    left_para.text = f'{case_name} [{case_id}]'
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_para.runs[0].font.size = Pt(9)
    
    # Center cell - Page number (witness number)
    center_cell = table.rows[0].cells[1]
    center_para = center_cell.paragraphs[0]
    center_para.text = str(witness_number)
    center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center_para.runs[0].font.size = Pt(9)
    center_para.runs[0].font.bold = True
    
    # Right cell - First line of confidentiality
    right_cell = table.rows[0].cells[2]
    right_para = right_cell.paragraphs[0]
    right_para.text = 'ATTORNEY WORK PRODUCT'
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.runs[0].font.size = Pt(9)
    right_para.runs[0].font.bold = True
    
    # Row 2: Date (left), Empty (center), Second line of confidentiality (right)
    # Left cell - Date
    left_cell_2 = table.rows[1].cells[0]
    left_para_2 = left_cell_2.paragraphs[0]
    left_para_2.text = footer_date_formatted
    left_para_2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_para_2.runs[0].font.size = Pt(9)
    
    # Right cell - Second line of confidentiality
    right_cell_2 = table.rows[1].cells[2]
    right_para_2 = right_cell_2.paragraphs[0]
    right_para_2.text = 'PRIVILEGED & CONFIDENTIAL'
    right_para_2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para_2.runs[0].font.size = Pt(9)
    right_para_2.runs[0].font.bold = True
    
    # Remove table borders and reduce spacing between rows
    table.autofit = False
    for idx, row in enumerate(table.rows):
        # Reduce row height
        tr = row._element
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '180')  # Reduced height in twips (180 = ~0.125 inches)
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)
        
        for cell in row.cells:
            # Remove borders
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            
            # Set all borders to none
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tcBorders.append(border)
            
            tcPr.append(tcBorders)
            
            # Reduce paragraph spacing within cells
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0

if __name__ == '__main__':
    app.run(debug=True, port=5000)